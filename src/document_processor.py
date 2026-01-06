"""
TKLocalAI - Document Processor
===============================
Handles document loading, parsing, and chunking for the RAG pipeline.
Supports PDFs, text files, markdown, and source code.
"""

import hashlib
import mimetypes
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
from datetime import datetime
from loguru import logger

# Document processing imports
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF not installed. PDF support limited.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX support disabled.")

import markdown
from bs4 import BeautifulSoup


@dataclass
class DocumentChunk:
    """Represents a chunk of a document."""
    id: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def __post_init__(self):
        if not self.id:
            # Generate ID from content hash
            self.id = hashlib.md5(self.content.encode()).hexdigest()[:12]


@dataclass  
class Document:
    """Represents a loaded document."""
    id: str
    path: str
    filename: str
    content: str
    file_type: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunks: List[DocumentChunk] = field(default_factory=list)
    
    def __post_init__(self):
        if not self.id:
            self.id = hashlib.md5(f"{self.path}:{self.filename}".encode()).hexdigest()[:16]


class DocumentLoader:
    """Loads documents from various file formats."""
    
    # Supported file extensions and their handlers
    SUPPORTED_EXTENSIONS = {
        # Text files
        '.txt': 'text',
        '.md': 'markdown',
        '.markdown': 'markdown',
        
        # Documents
        '.pdf': 'pdf',
        '.docx': 'docx',
        
        # Source code
        '.py': 'code',
        '.js': 'code',
        '.ts': 'code',
        '.jsx': 'code',
        '.tsx': 'code',
        '.java': 'code',
        '.cpp': 'code',
        '.c': 'code',
        '.h': 'code',
        '.hpp': 'code',
        '.cs': 'code',
        '.go': 'code',
        '.rs': 'code',
        '.rb': 'code',
        '.php': 'code',
        '.swift': 'code',
        '.kt': 'code',
        '.scala': 'code',
        '.r': 'code',
        '.sql': 'code',
        '.sh': 'code',
        '.bash': 'code',
        '.ps1': 'code',
        '.yaml': 'code',
        '.yml': 'code',
        '.json': 'code',
        '.xml': 'code',
        '.html': 'code',
        '.css': 'code',
        '.scss': 'code',
        '.less': 'code',
        
        # Config files
        '.ini': 'text',
        '.cfg': 'text',
        '.conf': 'text',
        '.toml': 'text',
        '.env': 'text',
    }
    
    # Programming language detection
    CODE_LANGUAGES = {
        '.py': 'python',
        '.js': 'javascript',
        '.ts': 'typescript',
        '.jsx': 'javascript',
        '.tsx': 'typescript',
        '.java': 'java',
        '.cpp': 'cpp',
        '.c': 'c',
        '.h': 'c',
        '.hpp': 'cpp',
        '.cs': 'csharp',
        '.go': 'go',
        '.rs': 'rust',
        '.rb': 'ruby',
        '.php': 'php',
        '.swift': 'swift',
        '.kt': 'kotlin',
        '.scala': 'scala',
        '.r': 'r',
        '.sql': 'sql',
        '.sh': 'shell',
        '.bash': 'bash',
        '.ps1': 'powershell',
    }
    
    @classmethod
    def is_supported(cls, file_path: Path) -> bool:
        """Check if a file type is supported."""
        return file_path.suffix.lower() in cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    def load(cls, file_path: Path) -> Optional[Document]:
        """
        Load a document from file.
        
        Args:
            file_path: Path to the document file
        
        Returns:
            Document object or None if loading failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        suffix = file_path.suffix.lower()
        handler_type = cls.SUPPORTED_EXTENSIONS.get(suffix)
        
        if not handler_type:
            logger.warning(f"Unsupported file type: {suffix}")
            return None
        
        try:
            # Load content based on type
            if handler_type == 'pdf':
                content = cls._load_pdf(file_path)
            elif handler_type == 'docx':
                content = cls._load_docx(file_path)
            elif handler_type == 'markdown':
                content = cls._load_markdown(file_path)
            elif handler_type == 'code':
                content = cls._load_code(file_path)
            else:
                content = cls._load_text(file_path)
            
            if not content:
                logger.warning(f"No content extracted from: {file_path}")
                return None
            
            # Build metadata
            metadata = {
                'file_type': handler_type,
                'extension': suffix,
                'size_bytes': file_path.stat().st_size,
                'modified_at': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                'created_at': datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            }
            
            # Add language info for code files
            if handler_type == 'code' and suffix in cls.CODE_LANGUAGES:
                metadata['language'] = cls.CODE_LANGUAGES[suffix]
            
            doc = Document(
                id="",  # Will be auto-generated
                path=str(file_path.parent),
                filename=file_path.name,
                content=content,
                file_type=handler_type,
                metadata=metadata,
            )
            
            logger.info(f"Loaded document: {file_path.name} ({len(content)} chars)")
            return doc
            
        except Exception as e:
            logger.error(f"Error loading {file_path}: {e}")
            return None
    
    @classmethod
    def _load_text(cls, file_path: Path) -> str:
        """Load plain text file."""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # Fallback with error handling
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @classmethod
    def _load_markdown(cls, file_path: Path) -> str:
        """Load and convert markdown to plain text."""
        raw_content = cls._load_text(file_path)
        
        # Convert markdown to HTML, then extract text
        html = markdown.markdown(raw_content)
        soup = BeautifulSoup(html, 'html.parser')
        
        # Get text while preserving some structure
        text = soup.get_text(separator='\n')
        
        # Clean up excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    @classmethod
    def _load_pdf(cls, file_path: Path) -> str:
        """Load PDF document."""
        if not PYMUPDF_AVAILABLE:
            logger.error("PyMuPDF not available for PDF processing")
            return ""
        
        text_parts = []
        
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
        
        return '\n\n'.join(text_parts)
    
    @classmethod
    def _load_docx(cls, file_path: Path) -> str:
        """Load Word document."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for DOCX processing")
            return ""
        
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        return '\n\n'.join(paragraphs)
    
    @classmethod
    def _load_code(cls, file_path: Path) -> str:
        """Load source code file with metadata."""
        content = cls._load_text(file_path)
        
        # Add file header for context
        suffix = file_path.suffix.lower()
        language = cls.CODE_LANGUAGES.get(suffix, 'unknown')
        
        header = f"# File: {file_path.name}\n# Language: {language}\n\n"
        
        return header + content


class TextChunker:
    """Splits documents into chunks for embedding."""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: Optional[List[str]] = None,
    ):
        """
        Initialize the chunker.
        
        Args:
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
            separators: List of separators to split on (in order of priority)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", ". ", " "]
    
    def chunk_document(self, document: Document) -> Document:
        """
        Split a document into chunks.
        
        Args:
            document: Document to chunk
        
        Returns:
            Document with chunks populated
        """
        chunks = self._split_text(document.content)
        
        document.chunks = []
        for i, chunk_text in enumerate(chunks):
            chunk = DocumentChunk(
                id=f"{document.id}_chunk_{i}",
                content=chunk_text,
                metadata={
                    **document.metadata,
                    'document_id': document.id,
                    'document_filename': document.filename,
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                }
            )
            document.chunks.append(chunk)
        
        logger.debug(f"Created {len(document.chunks)} chunks from {document.filename}")
        return document
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks using recursive separator splitting."""
        if len(text) <= self.chunk_size:
            return [text] if text.strip() else []
        
        return self._recursive_split(text, self.separators)
    
    def _recursive_split(self, text: str, separators: List[str]) -> List[str]:
        """Recursively split text using separators."""
        if not separators:
            # No more separators, force split by character
            return self._force_split(text)
        
        separator = separators[0]
        remaining_separators = separators[1:]
        
        splits = text.split(separator)
        
        chunks = []
        current_chunk = ""
        
        for split in splits:
            # Check if adding this split would exceed chunk size
            test_chunk = current_chunk + separator + split if current_chunk else split
            
            if len(test_chunk) <= self.chunk_size:
                current_chunk = test_chunk
            else:
                # Current chunk is full
                if current_chunk:
                    chunks.append(current_chunk)
                
                # Check if split itself needs to be broken down
                if len(split) > self.chunk_size:
                    # Recursively split with remaining separators
                    sub_chunks = self._recursive_split(split, remaining_separators)
                    
                    # Add sub-chunks with overlap handling
                    for sub_chunk in sub_chunks:
                        chunks.append(sub_chunk)
                    
                    current_chunk = ""
                else:
                    current_chunk = split
        
        # Don't forget the last chunk
        if current_chunk:
            chunks.append(current_chunk)
        
        # Apply overlap
        return self._apply_overlap(chunks)
    
    def _force_split(self, text: str) -> List[str]:
        """Force split by character count when no separators work."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunks.append(text[start:end])
            start = end - self.chunk_overlap
        
        return chunks
    
    def _apply_overlap(self, chunks: List[str]) -> List[str]:
        """Apply overlap between chunks."""
        if len(chunks) <= 1 or self.chunk_overlap <= 0:
            return chunks
        
        overlapped = [chunks[0]]
        
        for i in range(1, len(chunks)):
            prev_chunk = chunks[i - 1]
            current_chunk = chunks[i]
            
            # Get overlap from end of previous chunk
            overlap_text = prev_chunk[-self.chunk_overlap:] if len(prev_chunk) > self.chunk_overlap else prev_chunk
            
            # Find a good break point in the overlap
            for sep in ['\n\n', '\n', '. ', ' ']:
                if sep in overlap_text:
                    # Start overlap from the separator
                    idx = overlap_text.rfind(sep)
                    overlap_text = overlap_text[idx + len(sep):]
                    break
            
            # Combine overlap with current chunk
            overlapped_chunk = overlap_text + current_chunk
            overlapped.append(overlapped_chunk)
        
        return overlapped


class DocumentProcessor:
    """High-level document processing pipeline."""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: Optional[List[str]] = None,
    ):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Target chunk size
            chunk_overlap: Overlap between chunks
            separators: Separators for chunking
        """
        self.chunker = TextChunker(chunk_size, chunk_overlap, separators)
    
    def process_file(self, file_path: Path) -> Optional[Document]:
        """
        Process a single file.
        
        Args:
            file_path: Path to the file
        
        Returns:
            Processed document with chunks, or None if failed
        """
        document = DocumentLoader.load(file_path)
        
        if document:
            document = self.chunker.chunk_document(document)
        
        return document
    
    def process_directory(
        self,
        directory: Path,
        recursive: bool = True,
        extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """
        Process all documents in a directory.
        
        Args:
            directory: Directory path
            recursive: Whether to process subdirectories
            extensions: Optional filter for file extensions
        
        Returns:
            List of processed documents
        """
        directory = Path(directory)
        
        if not directory.exists():
            logger.error(f"Directory not found: {directory}")
            return []
        
        # Get all files
        if recursive:
            files = list(directory.rglob('*'))
        else:
            files = list(directory.glob('*'))
        
        # Filter to files only
        files = [f for f in files if f.is_file()]
        
        # Filter by extension if specified
        if extensions:
            extensions = [ext.lower() if ext.startswith('.') else f'.{ext.lower()}' for ext in extensions]
            files = [f for f in files if f.suffix.lower() in extensions]
        else:
            # Filter to supported extensions
            files = [f for f in files if DocumentLoader.is_supported(f)]
        
        # Process each file
        documents = []
        for file_path in files:
            doc = self.process_file(file_path)
            if doc:
                documents.append(doc)
        
        logger.info(f"Processed {len(documents)} documents from {directory}")
        return documents
    
    def get_all_chunks(self, documents: List[Document]) -> List[DocumentChunk]:
        """Extract all chunks from a list of documents."""
        chunks = []
        for doc in documents:
            chunks.extend(doc.chunks)
        return chunks
